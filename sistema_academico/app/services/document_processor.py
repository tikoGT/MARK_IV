# app/services/document_processor.py
import os
import re
import PyPDF2
from io import BytesIO
from typing import Dict, List, Tuple
import docx

class DocumentProcessor:
    """
    Clase para procesar documentos PDF y DOCX, extraer su contenido
    y estructurarlo para su uso en la generación de exámenes.
    """
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """
        Extrae texto de un archivo PDF
        """
        try:
            text = ""
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
            return text
        except Exception as e:
            print(f"Error al procesar PDF: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extrae texto de un archivo DOCX
        """
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # También extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return '\n'.join(full_text)
        except Exception as e:
            print(f"Error al procesar DOCX: {e}")
            return ""
    
    @staticmethod
    def process_document(file_path: str) -> Dict:
        """
        Procesa un documento basado en su extensión
        y retorna un diccionario con su contenido estructurado
        """
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            raw_text = DocumentProcessor.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            raw_text = DocumentProcessor.extract_text_from_docx(file_path)
        else:
            # Para archivos de texto plano
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    raw_text = file.read()
            except:
                # Intentar con otra codificación
                with open(file_path, 'r', encoding='latin-1') as file:
                    raw_text = file.read()
        
        # Procesar el texto para estructurarlo
        processed_content = DocumentProcessor.structure_content(raw_text)
        
        return {
            'file_path': file_path,
            'raw_text': raw_text,
            'processed_content': processed_content
        }
    
    @staticmethod
    def structure_content(text: str) -> Dict:
        """
        Estructura el contenido crudo del documento para facilitar 
        la generación de preguntas. Identifica secciones, definiciones, etc.
        """
        # Dividir por párrafos no vacíos
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        
        # Identificar títulos y secciones
        sections = {}
        current_section = "general"
        sections[current_section] = []
        
        for paragraph in paragraphs:
            # Detectar si es un título (heurística simple)
            is_title = False
            
            # Si es corto, tiene menos de 7 palabras y termina sin punto
            if len(paragraph) < 100 and len(paragraph.split()) < 7 and not paragraph.strip().endswith('.'):
                is_title = True
            
            # Si tiene números de sección como "1.", "1.1", etc.
            if re.match(r'^\d+(\.\d+)*\.?\s+\w+', paragraph):
                is_title = True
            
            # Si está todo en mayúsculas
            if paragraph.isupper() and len(paragraph) > 5:
                is_title = True
            
            if is_title:
                current_section = paragraph
                sections[current_section] = []
            else:
                sections[current_section].append(paragraph)
        
        # Identificar definiciones (patrones de texto como "X es Y", "X se define como Y")
        definitions = {}
        for section, paragraphs in sections.items():
            for paragraph in paragraphs:
                # Buscar patrones como "X es Y", "X se define como Y"
                matches = re.findall(r'([^.!?]+)(?:\ses\s|\sse\sdefine\scomo\s|\sse\srefiere\sa\s)([^.!?]+)[.!?]', paragraph)
                for match in matches:
                    term = match[0].strip()
                    definition = match[1].strip()
                    if term and definition:
                        definitions[term] = definition
        
        # Identificar listas numeradas o con viñetas
        lists = []
        list_items = []
        in_list = False
        
        for paragraph in paragraphs:
            # Detectar elementos de lista numerada o con viñetas
            if re.match(r'^\s*(\d+\.|•|\*|\-)\s', paragraph):
                if not in_list:
                    in_list = True
                    list_items = []
                list_items.append(paragraph)
            elif in_list:
                in_list = False
                if list_items:
                    lists.append(list_items.copy())
        
        if in_list and list_items:
            lists.append(list_items)
        
        return {
            'sections': sections,
            'definitions': definitions,
            'lists': lists,
            'total_paragraphs': len(paragraphs)
        }
    
    @staticmethod
    def find_concepts_for_questions(processed_content: Dict) -> List[Dict]:
        """
        Identifica conceptos que pueden convertirse en preguntas
        """
        concepts = []
        
        # Extraer de definiciones
        for term, definition in processed_content['definitions'].items():
            concepts.append({
                'type': 'definition',
                'term': term,
                'content': definition,
                'difficulty': 'medium'
            })
        
        # Extraer de secciones
        for section, paragraphs in processed_content['sections'].items():
            if section != "general" and paragraphs:
                # Tomar el primer párrafo de cada sección como explicación
                concepts.append({
                    'type': 'section_content',
                    'term': section,
                    'content': paragraphs[0] if paragraphs else "",
                    'difficulty': 'medium'
                })
        
        # Extraer de listas
        for list_items in processed_content['lists']:
            if len(list_items) >= 3:  # Solo listas con al menos 3 elementos
                list_title = "Lista"
                list_content = "\n".join(list_items)
                concepts.append({
                    'type': 'list_items',
                    'term': list_title,
                    'content': list_content,
                    'difficulty': 'easy'
                })
        
        return concepts