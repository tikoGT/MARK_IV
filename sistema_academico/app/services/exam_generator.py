# app/services/exam_generator.py
import os
import random
import re
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.services.document_processor import DocumentProcessor

class ExamGenerator:
    """
    Clase para generar exámenes a partir de contenido procesado.
    """
    
    @staticmethod
    def generate_questions_from_concept(concept: Dict, num_options: int = 4) -> List[Dict]:
        """
        Genera diferentes tipos de preguntas a partir de un concepto
        """
        questions = []
        question_type = concept['type']
        term = concept['term']
        content = concept['content']
        difficulty = concept['difficulty']
        
        if question_type == 'definition':
            # Pregunta "¿Qué es X?"
            questions.append({
                'content': f'¿Qué es {term}?',
                'question_type': 'open',
                'difficulty': difficulty,
                'points': 5.0,
                'answer': content
            })
            
            # Pregunta de selección múltiple
            questions.append({
                'content': f'¿Cuál de las siguientes opciones describe correctamente {term}?',
                'question_type': 'multiple_choice',
                'difficulty': difficulty,
                'points': 3.0,
                'options': ExamGenerator._generate_options(content, num_options),
                'answer': content
            })
        
        elif question_type == 'section_content':
            # Pregunta "Explique el concepto de X"
            questions.append({
                'content': f'Explique el concepto de {term}',
                'question_type': 'open',
                'difficulty': difficulty,
                'points': 5.0,
                'answer': content
            })
            
            # Pregunta "Verdadero o Falso"
            is_true = random.choice([True, False])
            if is_true:
                statement = content
                answer = "Verdadero"
            else:
                # Modificar el contenido para hacerlo falso
                words = content.split()
                if len(words) > 5:
                    pos = random.randint(2, len(words) - 3)
                    words[pos] = "NO" + words[pos]
                statement = " ".join(words)
                answer = "Falso"
            
            questions.append({
                'content': f'Indique si la siguiente afirmación es verdadera o falsa: "{statement}"',
                'question_type': 'multiple_choice',
                'difficulty': 'easy',
                'points': 2.0,
                'options': ['Verdadero', 'Falso'],
                'answer': answer
            })
        
        elif question_type == 'list_items':
            # Extraer elementos de la lista
            list_items = [item.strip() for item in content.split('\n')]
            list_items = [re.sub(r'^\s*(\d+\.|•|\*|\-)\s', '', item).strip() for item in list_items]
            
            if len(list_items) >= 4:
                # Pregunta "¿Cuál de los siguientes NO es parte de X?"
                valid_items = random.sample(list_items, 3)
                
                # Generar una opción incorrecta
                fake_item = f"Elemento que no es parte de {term}"
                
                options = valid_items + [fake_item]
                random.shuffle(options)
                
                questions.append({
                    'content': f'¿Cuál de los siguientes NO es parte de {term}?',
                    'question_type': 'multiple_choice',
                    'difficulty': 'medium',
                    'points': 3.0,
                    'options': options,
                    'answer': fake_item
                })
                
                # Pregunta "Enumere al menos 3 elementos de X"
                questions.append({
                    'content': f'Enumere al menos 3 elementos que forman parte de {term}',
                    'question_type': 'open',
                    'difficulty': 'medium',
                    'points': 4.0,
                    'answer': ", ".join(list_items[:3])
                })
        
        return questions
    
    @staticmethod
    def _generate_options(correct_answer: str, num_options: int) -> List[str]:
        """
        Genera opciones para preguntas de selección múltiple
        """
        options = [correct_answer]
        
        # Generar opciones incorrectas modificando la respuesta correcta
        words = correct_answer.split()
        
        for _ in range(num_options - 1):
            if len(words) > 4:
                # Modificar algunas palabras
                fake_answer = words.copy()
                num_changes = random.randint(1, min(3, len(words) // 3))
                
                for _ in range(num_changes):
                    pos = random.randint(0, len(fake_answer) - 1)
                    # Reemplazar, eliminar o invertir palabras
                    action = random.choice(['replace', 'remove', 'swap'])
                    
                    if action == 'replace':
                        fake_answer[pos] = f"NO-{fake_answer[pos]}"
                    elif action == 'remove':
                        fake_answer[pos] = ""
                    elif action == 'swap' and pos < len(fake_answer) - 1:
                        fake_answer[pos], fake_answer[pos+1] = fake_answer[pos+1], fake_answer[pos]
                
                fake_text = " ".join([w for w in fake_answer if w])
                if fake_text and fake_text != correct_answer and fake_text not in options:
                    options.append(fake_text)
                else:
                    # Si no se pudo generar una opción válida, crear una genérica
                    options.append(f"Opción incorrecta {len(options)}")
            else:
                # Para respuestas cortas, crear opciones genéricas
                options.append(f"Opción incorrecta {len(options)}")
        
        random.shuffle(options)
        return options
    
    @staticmethod
    def create_exam_from_materials(
        materials: List[Dict], 
        exam_config: Dict
    ) -> Dict:
        """
        Crea un examen a partir de materiales procesados
        """
        # Extraer conceptos de todos los materiales
        all_concepts = []
        for material in materials:
            processed_content = material.get('processed_content', {})
            concepts = DocumentProcessor.find_concepts_for_questions(processed_content)
            all_concepts.extend(concepts)
        
        # Seleccionar conceptos según la configuración
        selected_concepts = ExamGenerator._select_concepts(
            all_concepts, 
            exam_config.get('num_concepts', 10),
            exam_config.get('difficulty_distribution', {'easy': 0.3, 'medium': 0.5, 'hard': 0.2})
        )
        
        # Generar preguntas a partir de los conceptos seleccionados
        questions = []
        for concept in selected_concepts:
            concept_questions = ExamGenerator.generate_questions_from_concept(
                concept, 
                num_options=exam_config.get('num_options', 4)
            )
            questions.extend(concept_questions)
        
        # Seleccionar preguntas según la configuración
        selected_questions = ExamGenerator._select_questions(
            questions, 
            exam_config.get('num_questions', 20),
            exam_config.get('question_type_distribution', {
                'multiple_choice': 0.7, 
                'open': 0.3
            })
        )
        
        return {
            'title': exam_config.get('title', 'Examen Generado'),
            'description': exam_config.get('description', 'Examen generado automáticamente'),
            'questions': selected_questions,
            'total_points': sum(q['points'] for q in selected_questions)
        }
    
    @staticmethod
    def _select_concepts(
        concepts: List[Dict], 
        num_concepts: int,
        difficulty_distribution: Dict[str, float]
    ) -> List[Dict]:
        """
        Selecciona conceptos según la distribución de dificultad deseada
        """
        if not concepts:
            return []
        
        # Agrupar conceptos por dificultad
        by_difficulty = {'easy': [], 'medium': [], 'hard': []}
        
        for concept in concepts:
            difficulty = concept.get('difficulty', 'medium')
            by_difficulty[difficulty].append(concept)
        
        # Calcular cuántos conceptos de cada dificultad seleccionar
        counts = {}
        remaining = num_concepts
        
        for difficulty, percentage in difficulty_distribution.items():
            count = int(num_concepts * percentage)
            counts[difficulty] = min(count, len(by_difficulty[difficulty]))
            remaining -= counts[difficulty]
        
        # Distribuir conceptos restantes
        while remaining > 0:
            for difficulty in ['easy', 'medium', 'hard']:
                if remaining > 0 and len(by_difficulty[difficulty]) > counts[difficulty]:
                    counts[difficulty] += 1
                    remaining -= 1
        
        # Seleccionar conceptos
        selected = []
        for difficulty, count in counts.items():
            selected.extend(random.sample(by_difficulty[difficulty], min(count, len(by_difficulty[difficulty]))))
        
        return selected
    
    @staticmethod
    def _select_questions(
        questions: List[Dict], 
        num_questions: int,
        question_type_distribution: Dict[str, float]
    ) -> List[Dict]:
        """
        Selecciona preguntas según la distribución de tipos deseada
        """
        if not questions:
            return []
        
        # Agrupar preguntas por tipo
        by_type = {}
        for question in questions:
            q_type = question.get('question_type', 'multiple_choice')
            if q_type not in by_type:
                by_type[q_type] = []
            by_type[q_type].append(question)
        
        # Calcular cuántas preguntas de cada tipo seleccionar
        counts = {}
        remaining = num_questions
        
        for q_type, percentage in question_type_distribution.items():
            if q_type in by_type:
                count = int(num_questions * percentage)
                counts[q_type] = min(count, len(by_type[q_type]))
                remaining -= counts[q_type]
        
        # Distribuir preguntas restantes
        while remaining > 0:
            for q_type in by_type.keys():
                if remaining > 0 and q_type in counts and len(by_type[q_type]) > counts[q_type]:
                    counts[q_type] += 1
                    remaining -= 1
        
        # Seleccionar preguntas
        selected = []
        for q_type, count in counts.items():
            if q_type in by_type:
                selected.extend(random.sample(by_type[q_type], min(count, len(by_type[q_type]))))
        
        # Ordenar por tipo de pregunta (primero opción múltiple, luego abiertas)
        selected.sort(key=lambda q: 0 if q.get('question_type') == 'multiple_choice' else 1)
        
        return selected
    
    @staticmethod
    def generate_exam_docx(exam: Dict, output_path: str, include_answers: bool = False) -> str:
        """
        Genera un documento DOCX con el examen
        """
        doc = Document()
        
        # Estilo del título
        title = doc.add_heading(exam['title'], level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Descripción
        if exam.get('description'):
            description = doc.add_paragraph(exam['description'])
            description.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información del examen
        info = doc.add_paragraph()
        info.add_run('Total de puntos: ').bold = True
        info.add_run(f"{exam['total_points']}")
        
        # Agregar línea divisoria
        doc.add_paragraph('_' * 50)
        
        # Instrucciones
        doc.add_heading('Instrucciones:', level=2)
        doc.add_paragraph('Responda cada una de las siguientes preguntas. Lea cuidadosamente antes de responder.')
        
        # Preguntas
        questions = exam['questions']
        
        for i, question in enumerate(questions, 1):
            # Título de la pregunta
            q_title = doc.add_paragraph()
            q_title.add_run(f"Pregunta {i} ").bold = True
            q_title.add_run(f"({question['points']} puntos) ").italic = True
            q_title.add_run(f"[{question['difficulty']}] ").italic = True
            
            # Contenido de la pregunta
            doc.add_paragraph(question['content'])
            
            # Opciones para preguntas de selección múltiple
            if question['question_type'] == 'multiple_choice' and 'options' in question:
                for j, option in enumerate(question['options']):
                    option_text = f"{chr(65 + j)}) {option}"
                    option_para = doc.add_paragraph(option_text, style='List Bullet')
                    
                    # Si se incluyen respuestas, marcar la correcta
                    if include_answers and option == question.get('answer'):
                        for run in option_para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0, 128, 0)  # Verde
            
            # Espacio para respuestas abiertas
            if question['question_type'] == 'open':
                doc.add_paragraph("Respuesta:", style='List Bullet')
                for _ in range(4):  # Agregar líneas para respuesta
                    doc.add_paragraph("_" * 50)
                
                # Si se incluyen respuestas, mostrar la correcta
                if include_answers and 'answer' in question:
                    answer_para = doc.add_paragraph()
                    answer_para.add_run("Respuesta modelo: ").bold = True
                    answer_para.add_run(question['answer'])
                    for run in answer_para.runs:
                        run.font.color.rgb = RGBColor(0, 128, 0)  # Verde
            
            # Separador entre preguntas
            doc.add_paragraph()
        
        # Guardar el documento
        doc.save(output_path)
        
        return output_path
    
    @staticmethod
    def generate_answer_sheet(exam: Dict, output_path: str) -> str:
        """
        Genera una hoja de respuestas para el examen
        """
        doc = Document()
        
        # Estilo del título
        title = doc.add_heading(f"Hoja de Respuestas: {exam['title']}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información del estudiante
        doc.add_paragraph('Nombre: ______________________________')
        doc.add_paragraph('ID: ______________________________')
        doc.add_paragraph('Fecha: ______________________________')
        
        # Agregar línea divisoria
        doc.add_paragraph('_' * 50)
        
        # Crear tabla para respuestas de selección múltiple
        multiple_choice_questions = [q for q in exam['questions'] if q['question_type'] == 'multiple_choice']
        
        if multiple_choice_questions:
            doc.add_heading('Selección Múltiple', level=2)
            
            # Organizamos 10 preguntas por fila, con opciones A, B, C, D por columna
            rows_needed = (len(multiple_choice_questions) + 9) // 10
            
            for row in range(rows_needed):
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # Encabezado
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Pregunta'
                hdr_cells[1].text = 'A'
                hdr_cells[2].text = 'B'
                hdr_cells[3].text = 'C'
                hdr_cells[4].text = 'D'
                
                # Filas de preguntas
                for i in range(10):
                    q_idx = row * 10 + i
                    if q_idx < len(multiple_choice_questions):
                        q_num = exam['questions'].index(multiple_choice_questions[q_idx]) + 1
                        
                        row_cells = table.add_row().cells
                        row_cells[0].text = f"{q_num}"
                        
                        # Círculos para marcar
                        for j in range(1, 5):
                            # Si la pregunta tiene menos opciones, no mostramos el círculo
                            if j-1 < len(multiple_choice_questions[q_idx].get('options', [])):
                                row_cells[j].text = "○"
                            else:
                                row_cells[j].text = ""
                
                doc.add_paragraph()
        
        # Espacio para preguntas abiertas
        open_questions = [q for q in exam['questions'] if q['question_type'] == 'open']
        
        if open_questions:
            doc.add_heading('Preguntas Abiertas', level=2)
            
            for question in open_questions:
                q_num = exam['questions'].index(question) + 1
                doc.add_paragraph(f"Pregunta {q_num}:")
                
                # Agregar líneas para la respuesta
                for _ in range(8):
                    doc.add_paragraph("_" * 60)
                
                doc.add_paragraph()
        
        # Guardar el documento
        doc.save(output_path)
        
        return output_path